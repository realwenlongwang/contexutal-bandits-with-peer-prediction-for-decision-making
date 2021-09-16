from docx import Document
from docx.shared import Cm
import matplotlib.pyplot as plt
import numpy as np
from Environment import *
from tqdm import tqdm, trange
from PolicyGradientAgent import StochasticGradientAgent
import traceback
from sys import exit
from scipy.ndimage import uniform_filter1d
from scipy.special import logit, expit
import pandas as pd

root_dir = 'plots/'

def rewards_fig(reward_history_df, file_name):
    fig, axs = plt.subplots(3, figsize=(15, 8))
    axs[0].scatter(x=reward_history_df.index, y=reward_history_df['actual_reward'], label='Actual log rewards',
                   marker='.', s=3)
    axs[1].plot(no_outlier_array(reward_history_df.iloc[:, 1]), 'y', zorder=-99, label='Average reward')
    axs[2].plot(reward_history_df['max_expected_reward'] - reward_history_df['expected_reward'], 'g.', label='Regret')
    axs[0].hlines(y=0.0, xmin=0, xmax=reward_history_df.shape[0], colors='black', linestyles='dashdot')
    axs[1].hlines(y=0.0, xmin=0, xmax=reward_history_df.shape[0], colors='black', linestyles='dashdot')
    for signal, df in reward_history_df.reset_index().groupby('prior_red'):
        axs[1].scatter(x=df['index'], y=df['estimated_reward'], label='Estimated reward ' + str(signal), marker='.',
                       s=3)
    fig.legend()
    fig.suptitle('Actual Rewards and Average')
    img_dir = root_dir + file_name + '_reward.png'
    plt.savefig(img_dir, dpi=150)
    return img_dir

def report_fig(report_history_df, prior_red_list, pr_red_ball_red_bucket, pr_red_ball_blue_bucket,file_name):
    fig, ax = plt.subplots(figsize=(15, 4))
    for signal, df in report_history_df.reset_index().groupby('signal'):
        ax.scatter(x=df['index'], y=df['report'], label=signal, marker='.', c=signal, s=3, zorder=-99)
    for prior_red in prior_red_list:
        plt.hlines(
            y=analytical_best_report_ru_rs(
                pr_ru=prior_red,
                pr_rs_ru=pr_red_ball_red_bucket,
                pr_rs_bu=pr_red_ball_blue_bucket
            ), xmin=0, xmax=len(report_history_df), colors='black', linestyles='dashdot', zorder=-98)

        plt.hlines(
            y=analytical_best_report_ru_bs(
                pr_ru=prior_red,
                pr_bs_ru=1 - pr_red_ball_red_bucket,
                pr_bs_bu=1 - pr_red_ball_blue_bucket
            ), xmin=0, xmax=len(report_history_df), colors='black', linestyles='dashdot', zorder=-98)
    ax.legend(loc='lower left')
    plt.title('Report')
    img_dir = root_dir + file_name + '_report.png'
    plt.savefig(img_dir, dpi=150)
    return img_dir


def weights_for_mean_fig(mean_weights_history_df, file_name, pr_red_ball_red_bucket, pr_red_ball_blue_bucket):
    fig = plt.figure(figsize=(15, 4))
    plt.plot(mean_weights_history_df.iloc[1:, 0], 'r', label='Red weight')
    plt.plot(mean_weights_history_df.iloc[1:, 1], label='Blue weight')
    plt.plot(mean_weights_history_df.iloc[1:, 2], 'g', label='Prior weight')
    plt.hlines(y=logit(pr_red_ball_red_bucket), xmin=0, xmax=len(mean_weights_history_df), colors='red',
               linestyles='dashdot')
    plt.annotate('%.3f' % logit(pr_red_ball_red_bucket),
                 xy=(len(mean_weights_history_df) / 2, logit(pr_red_ball_red_bucket)),
                 xytext=(len(mean_weights_history_df) / 2, np.log(2) / 2), arrowprops=dict(arrowstyle="->"))
    plt.hlines(y=logit(pr_red_ball_blue_bucket), xmin=0, xmax=len(mean_weights_history_df), colors='blue',
               linestyles='dashdot')
    plt.annotate('%.3f' % logit(pr_red_ball_blue_bucket),
                 xy=(len(mean_weights_history_df) / 2, logit(pr_red_ball_blue_bucket)),
                 xytext=(len(mean_weights_history_df) / 2, np.log(1 / 2) / 2), arrowprops=dict(arrowstyle="->"))
    plt.hlines(y=1, xmin=0, xmax=len(mean_weights_history_df), colors='green', linestyles='dashdot')
    # for coord in phase_change_coordinates(mark_index, mean_weights_history_df):
    #     plt.annotate('change', xy=coord, xytext=(coord[0], 0.1), arrowprops=dict(arrowstyle="->"))
    plt.legend()
    plt.title('Weights for Mean')

    img_dir = root_dir + file_name + '_weights.png'
    plt.savefig(img_dir, dpi=150)
    return img_dir

def gradients_for_mean_fig(grad_mean_history_df, file_name):
    fig, ax = plt.subplots(figsize=(15, 4))
    grad_mean_history_df.iloc[100:, :].plot(ax=ax, color=['red', 'blue', 'green'], zorder=-100)
    ax.hlines(y=0, xmin=0, xmax=len(grad_mean_history_df), linestyles='dashdot', zorder=-99)
    img_dir = root_dir + file_name + '_gradients.png'
    plt.savefig(img_dir, dpi=150)
    return img_dir

def successive_gradients_dot_product(grad_mean_history_df, file_name, moving_size = 1000):
    grad_mean_successive_dot = np.sum(grad_mean_history_df.values * np.roll(grad_mean_history_df.values, 1, axis=0), axis=1)[1:]
    moving_size = 1000
    fig, axs = plt.subplots(2, figsize=(15, 8))
    axs[0].plot(grad_mean_successive_dot[100:], zorder=-100)
    axs[0].hlines(y=0, xmin=0, xmax=len(grad_mean_successive_dot), linestyles='dashdot', color='black', zorder=-99)
    axs[0].set_title('Successive gradients dot product')
    axs[1].plot(uniform_filter1d(grad_mean_successive_dot[100:], size=moving_size), zorder=-100)
    axs[1].hlines(y=0, xmin=0, xmax=len(grad_mean_successive_dot), linestyles='dashdot', color='black', zorder=-99)
    axs[1].set_title('Successive gradients dot product size %i moving average'%moving_size)

    img_dir = root_dir + file_name + '_dot_product.png'
    plt.savefig(img_dir, dpi=150)
    return img_dir

def pd_table_to_fig(data, title, file_name, footer='', fig_background_color='skyblue', fig_border='steelblue'):
    rcolors = plt.cm.BuPu(np.full(len(data.index), 0.1))
    ccolors = plt.cm.BuPu(np.full(len(data.columns), 0.1))
    # Create the figure. Setting a small pad on tight_layout
    # seems to better regulate white space. Sometimes experimenting
    # with an explicit figsize here can produce better outcome.
    plt.figure(linewidth=2,
               edgecolor=fig_border,
               facecolor=fig_background_color,
               tight_layout={'pad':1},
               figsize=(10,6)
              )
    # Hide axes
    ax = plt.gca()
    ax.get_xaxis().set_visible(False)
    ax.get_yaxis().set_visible(False)

    the_table = pd.plotting.table(ax=ax, data=data,
                                  rowColours=rcolors,
                                  rowLoc='right',
                                  colColours=ccolors,
                                  loc='center')
    the_table.scale(1, 1.5)
    # Hide axes border
    plt.box(on=None)
    # Add title
    plt.suptitle(title)
    # Add footer
    plt.figtext(0.95, 0.05, footer, horizontalalignment='right', size=6, weight='light')
    # Force the figure to update, so backends center objects correctly within the figure.
    # Without plt.draw() here, the title will center on the axes and not the figure.
    plt.draw()
    # Create image. plt.savefig ignores figure edge and face colors, so map them.
    fig = plt.gcf()

    img_dir = root_dir + file_name + '_' + title + '.png'
    plt.savefig(img_dir, dpi=150)
    return img_dir

def main_loop(learning_rate_theta, learning_rate_wv, memory_size, training_episodes, fixed_std, algorithm, prior_red_list=[1/2, 1/2], pr_red_ball_red_bucket=2/3,  pr_red_ball_blue_bucket=1/3):
    # learning_rate_theta = learning_rate_theta
    # learning_rate_wv = learning_rate_wv
    # memory_size = memory_size
    batch_size = memory_size
    # training_episodes = training_episodes
    decay_rate = 0
    beta1 = 0.9
    beta2 = 0.9999
    # Algorithm: adam, momentum, regular
    # algorithm = 'regular'
    if fixed_std != 0:
        learning_std = False
    else:
        learning_std = True
    # Bucket parameters
    prior_red_list = [3/4, 1/4]
    pr_red_ball_red_bucket = 2/3
    pr_red_ball_blue_bucket = 1/3

    agent = StochasticGradientAgent(feature_shape=[1, 3], learning_rate_theta=learning_rate_theta,
                                    learning_rate_wv=learning_rate_wv,
                                    memory_size=memory_size, batch_size=batch_size,
                                    beta1=beta1, beta2=beta2,
                                    learning_std=learning_std, fixed_std=fixed_std)

    reward_history_list = []
    average_reward = 0

    mean_weights_history_list = []
    std_weights_history_list = []

    r_ball_mean_history_list = []
    b_ball_mean_history_list = []
    r_ball_std_history_list = []
    b_ball_std_history_list = []

    r_ball_pred_history_list = []
    b_ball_pred_history_list = []

    grad_r_ball_mean_history_list = []
    grad_b_ball_mean_history_list = []
    grad_r_ball_std_history_list = []
    grad_b_ball_std_history_list = []

    grad_r_ball_v_mean_history_list = []
    grad_b_ball_v_mean_history_list = []
    grad_r_ball_v_std_history_list = []
    grad_b_ball_v_std_history_list = []

    grad_r_ball_adam_mean_history_list = []
    grad_b_ball_adam_mean_history_list = []
    grad_r_ball_adam_std_history_list = []
    grad_b_ball_adam_std_history_list = []

    grad_mean_history_list = []
    mean_history_list = []
    std_history_list = []
    report_history_list = []

    for t in trange(training_episodes):
        prior_red = np.random.choice(prior_red_list)
        #     prior_red = np.random.uniform()
        bucket = Bucket(prior_red, pr_red_ball_red_bucket, pr_red_ball_blue_bucket)
        pm = PredictionMarket(prior_red=prior_red)
        signal = bucket.signal()
        x = one_hot_encode(signal)
        x.append(logit(prior_red))
        h, mean, std = agent.report(x)
        pi = expit(h)
        report = [pi, 1 - pi]

        pm.report(report)
        R = pm.log_resolve(bucket_colour_to_num[bucket.colour])

        average_reward = average_reward + (1 / (t + 1)) * (R - average_reward)

        mean_weights_history_list.append(agent.theta_mean[0].tolist())
        std_weights_history_list.append(agent.theta_std[0].tolist())

        actual_pr_ru_S = 0
        expected_log_reward = 0
        max_expected_log_reward = 0
        regret = 0
        if signal == 'red':
            actual_pr_ru_S = analytical_best_report_ru_rs(pr_ru=prior_red, pr_rs_ru=pr_red_ball_red_bucket,
                                                          pr_rs_bu=pr_red_ball_blue_bucket)
            expected_log_reward = expected_log_reward_red_ball(actual_pr_ru_rs=actual_pr_ru_S, estimated_pr_ru_rs=pi,
                                                               pr_ru=prior_red)
            max_expected_log_reward = expected_log_reward_red_ball(actual_pr_ru_rs=actual_pr_ru_S,
                                                                   estimated_pr_ru_rs=actual_pr_ru_S, pr_ru=prior_red)
        else:
            actual_pr_ru_S = analytical_best_report_ru_bs(pr_ru=prior_red, pr_bs_ru=1 - pr_red_ball_red_bucket,
                                                          pr_bs_bu=1 - pr_red_ball_blue_bucket)
            expected_log_reward = expected_log_reward_blue_ball(actual_pr_ru_bs=actual_pr_ru_S, estimated_pr_ru_bs=pi,
                                                                pr_ru=prior_red)
            max_expected_log_reward = expected_log_reward_blue_ball(actual_pr_ru_bs=actual_pr_ru_S,
                                                                    estimated_pr_ru_bs=actual_pr_ru_S, pr_ru=prior_red)

        v = agent.store_experience(x, h, mean, std, R, t)

        reward_history_list.append([R, average_reward, v, expected_log_reward, max_expected_log_reward, signal, prior_red])
        try:
            grad_mean, grad_std, v_dw_mean_corrected, v_dw_std_corrected, \
            s_dw_mean_corrected, s_dw_std_corrected = agent.batch_update(t, algorithm=algorithm)
        except AssertionError:
            tb = traceback.format_exc()
            print(tb)

        agent.learning_rate_decay(epoch=t, decay_rate=decay_rate)

        if signal == 'red':
            r_ball_pred_history_list.append(report[0])
            r_ball_mean_history_list.append(mean)
            r_ball_std_history_list.append(std)

        else:
            b_ball_pred_history_list.append(report[0])
            b_ball_mean_history_list.append(mean)
            b_ball_std_history_list.append(std)

        report_history_list.append([report[0], signal])
        mean_history_list.append([mean, signal])
        std_history_list.append([std, signal])
        grad_mean_history_list.append(grad_mean[0, :])

        grad_r_ball_mean_history_list.append(grad_mean[0, 0])
        grad_r_ball_std_history_list.append(grad_std[0, 0])
        ##########
        grad_r_ball_v_mean_history_list.append(v_dw_mean_corrected[0, 0])
        grad_r_ball_v_std_history_list.append(v_dw_std_corrected[0, 0])
        grad_r_ball_adam_mean_history_list.append(s_dw_mean_corrected[0, 0])
        grad_r_ball_adam_std_history_list.append(s_dw_std_corrected[0, 0])
        ##################################################################
        grad_b_ball_mean_history_list.append(grad_mean[0, 1])
        grad_b_ball_std_history_list.append(grad_std[0, 1])
        ##########
        grad_b_ball_v_mean_history_list.append(v_dw_mean_corrected[0, 1])
        grad_b_ball_v_std_history_list.append(v_dw_std_corrected[0, 1])
        grad_b_ball_adam_mean_history_list.append(s_dw_mean_corrected[0, 1])
        grad_b_ball_adam_std_history_list.append(s_dw_std_corrected[0, 1])

    return reward_history_list, report_history_list, mean_weights_history_list, grad_mean_history_list, std_history_list[-1][0]


def generating_report(document, learning_rate_space, learning_rate_wv_space ,memory_size_space, fixed_std_space, algorithm_space):

    global reward_history_list
    for learning_rate in learning_rate_space:
        for learning_rate_wv in learning_rate_wv_space:
            for ms in memory_size_space:
                for std in fixed_std_space:
                    for algorithm in algorithm_space:
                        if algorithm == 'adam':
                            lr = learning_rate/25
                            lr_wv = learning_rate_wv/35
                        else:
                            lr = learning_rate
                            lr_wv = learning_rate_wv

                        prior_red_list=[3/4, 1/4]
                        pr_red_ball_red_bucket = 2 / 3
                        pr_red_ball_blue_bucket = 1 / 3
                        learning_rate_theta_string = 'learning_rate_theta: ' + str(lr)
                        learning_rate_wv_string = 'learning_rate_wv: ' + str(lr_wv)
                        memory_size_string = 'memory_size: ' + str(ms)
                        if std == 0:
                            std_string = 'standard deviation: learnable'
                        else:
                            std_string = 'standard deviation: ' + str(std)

                        learning_rate_p = document.add_paragraph(learning_rate_theta_string)
                        learning_rate_wv_p = document.add_paragraph(learning_rate_wv_string)
                        memory_size_p = document.add_paragraph(memory_size_string)
                        std_p = document.add_paragraph(std_string)
                        algorithm_p = document.add_paragraph(algorithm)

                        filename = 'lr' + '%.0E'%lr + '_v' + '%.0E'%lr_wv + '_ms' + str(ms) + '_std' + str(std)
                        # learning_rate_theta, learning_rate_wv, memory_size, training_episodes, fixed_std

                        try:
                            reward_history_list, report_history_list, \
                            mean_weights_history_list, \
                            grad_mean_history_list, final_std = main_loop(
                                learning_rate_theta=lr,
                                learning_rate_wv=lr_wv,
                                memory_size=ms,
                                training_episodes=900000 * 2,
                                fixed_std=std,
                                algorithm=algorithm,
                                prior_red_list=prior_red_list,
                                pr_red_ball_red_bucket=pr_red_ball_red_bucket,
                                pr_red_ball_blue_bucket=pr_red_ball_blue_bucket
                            )
                        except AssertionError:
                            document.save('report.docx')
                            exit(1)

                        if std == 0:
                            std_p.add_run(' final std: ' + str(final_std))

                        reward_history_df = pd.DataFrame(reward_history_list, columns=['actual_reward', 'average_reward', 'estimated_reward', 'expected_reward', 'max_expected_reward','signal', 'prior_red'])
                        report_history_df = pd.DataFrame(report_history_list, columns=['report', 'signal'])
                        grad_mean_history_df = pd.DataFrame(grad_mean_history_list, columns=['red_ball', 'blue_ball', 'prior'])
                        mean_weights_history_df = pd.DataFrame(mean_weights_history_list, columns=['red_weight', 'blue_weight', 'prior_weight'])


                        reward_img_dir = rewards_fig(reward_history_df=reward_history_df, file_name=filename)
                        report_img_dir = report_fig(
                            report_history_df=report_history_df,
                            file_name=filename,
                            prior_red_list=prior_red_list,
                            pr_red_ball_red_bucket=pr_red_ball_red_bucket,
                            pr_red_ball_blue_bucket=pr_red_ball_blue_bucket
                        )
                        weights_img_dir = weights_for_mean_fig(
                            mean_weights_history_df=mean_weights_history_df,
                            file_name=filename,
                            pr_red_ball_red_bucket=pr_red_ball_red_bucket,
                            pr_red_ball_blue_bucket=pr_red_ball_blue_bucket
                        )
                        gradients_for_mean_img_dir = gradients_for_mean_fig(grad_mean_history_df=grad_mean_history_df,
                                                                            file_name=filename)
                        successive_grad_dot_product_img_dir = successive_gradients_dot_product(grad_mean_history_df=grad_mean_history_df,
                                                                                                file_name=filename,
                                                                                                moving_size=1000
                                                                                                )
                        query_index = mean_weights_history_df.shape[0]//3
                        mean_grad_summary_img_dir = pd_table_to_fig(data=grad_mean_history_df.iloc[-query_index:, :].describe(), title='mean_gradients_history_summary', file_name=filename)
                        mean_weights_summary_img_dir = pd_table_to_fig(data=mean_weights_history_df.iloc[-query_index:, :].describe(), title='mean_weights_history_summary', file_name=filename)

                        document.add_picture(reward_img_dir, width=Cm(15.0))
                        document.add_picture(report_img_dir, width=Cm(15.0))
                        document.add_picture(weights_img_dir, width=Cm(15.0))
                        document.add_picture(gradients_for_mean_img_dir, width=Cm(15.0))
                        document.add_picture(successive_grad_dot_product_img_dir, width=Cm(15.0))
                        document.add_picture(mean_grad_summary_img_dir, width=Cm(15.0))
                        document.add_picture(mean_weights_summary_img_dir, width=Cm(15.0))

    return document


if __name__ == '__main__':

    learning_rate_space = [3e-5, 1e-4, 3e-4] # 3e-3, 1e-2 seems too high for regular algorithm
    learning_rate_wv_space = [0]
    memory_size_space = [1024]
    std_space = [0.3, 0] # 0 means learning the standard deviation
    algorithm_space = ['regular', 'momentum', 'adam']
    # baseline comparison, training iteration dependent on learning rate.


    document = Document()


    document.add_heading('Stochastic Policy Gradient Hyper-parameter Report', 0)
    # document.add_heading('Expected reward', level=1)
    # document = generating_report(document=document, file_prefix='er', learning_rate_space=learning_rate_space, memory_size_space=memory_size_space)


    document.add_heading('Actual reward', level=1)
    document = generating_report(document=document, learning_rate_space=learning_rate_space,
                                 learning_rate_wv_space=learning_rate_wv_space,
                                 memory_size_space=memory_size_space, fixed_std_space=std_space,
                                 algorithm_space=algorithm_space
                                 )


    document.save('stochastic_report.docx')
